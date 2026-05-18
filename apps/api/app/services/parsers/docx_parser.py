from docx import Document


def parse_docx(file_path: str) -> tuple[str, str | None]:
    try:
        doc = Document(file_path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        full_text = "\n".join(parts)
        if not full_text.strip():
            return "", "DOCX 文件未提取到文本内容（可能是图片型文档）"
        return full_text.strip(), None
    except Exception as e:
        return "", f"DOCX 解析失败: {str(e)[:300]}"
